"""
Generates the tailored resume + cover letter for BrightEdge in both .docx and .pdf.
Single source of truth for content — both formats stay in sync.

Run: python3 scripts/generate_brightedge_docs.py
Outputs:
  resumes/tailored/brightedge_cloud-finops-pmo_2026-05-25.docx
  resumes/tailored/brightedge_cloud-finops-pmo_2026-05-25.pdf
  resumes/tailored/brightedge_cloud-finops-pmo_2026-05-25_cover.docx
  resumes/tailored/brightedge_cloud-finops-pmo_2026-05-25_cover.pdf
"""
from __future__ import annotations
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

OUT_DIR = Path(__file__).resolve().parent.parent / "resumes" / "tailored"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SLUG = "brightedge_cloud-finops-pmo_2026-05-25"

# ---------- CONTENT (single source of truth) ----------

CONTACT = {
    "name": "SRIKANTH CHINTALA",
    "headline": "PMO & Operations Analyst | Cost Reporting, Invoice Coordination & FinOps-Aspiring",
    "location": "Hyderabad, India 500088",
    "phone": "+91 81795 98723",
    "email": "srikanthchintala191997@gmail.com",
    "linkedin": "linkedin.com/in/chintalasrikanth97",
}

SUMMARY = (
    "PMO and operations professional with 4+ years at Deloitte driving end-to-end project governance, "
    "financial tracking, and reporting for global delivery engagements across UK and European clients. "
    "Built and ran the RAG/RAID, dashboard, invoice/PO/SOW tracking, and budget-vs-actual machinery that "
    "kept 10+ multi-stakeholder, time-zone-spanning engagements on time and audit-ready. Strongest in "
    "turning messy, fast-paced, SaaS-style delivery into clean, leadership-ready trackers — with a record "
    "of 100% on-time filings, 25% reporting-effort reduction, 15 hrs/week of follow-ups eliminated, and "
    "4 RFP wins. Actively building Cloud FinOps depth (FinOps Foundation self-study, AWS CLF-C02 planned). "
    "MBA Finance + CA Inter; immediate joiner."
)

COMPETENCIES = [
    ("Cost Operations & Reporting",
     "Invoice / PO / SOW / MSA Tracking | Vendor & Billing Coordination | Cost Anomaly Detection | "
     "Variance & Plan-vs-Actual Analysis | Forecasting & Extrapolation | Budget Monitoring | "
     "Trackers, Dashboards & Spreadsheets | Operational Reviews"),
    ("PMO & Governance",
     "Project Management Office (PMO) | Project Governance | RAG Status Reporting | RAID Log Management | "
     "Stakeholder Management | Scope, Deliverable & Timeline Management | Change Request Tracking | "
     "Audit Readiness | Documentation Control"),
    ("Project Delivery & Operations",
     "Project Planning & Scheduling | Resource & Capacity Planning | Risk Management | "
     "Cross-functional Collaboration | Process Improvement | Agile / Waterfall | Implementation Lifecycle | "
     "Performance Metrics & KPIs | Vendor Coordination | Onboarding Management"),
    ("Tools & Excel Depth",
     "Advanced MS Excel — Pivot Tables, XLOOKUP / INDEX-MATCH, SUMIFS, Power Query, Conditional "
     "Formatting, Data Validation, Basic Macros | Power BI | MS Project | MS PowerPoint | Jira | "
     "SharePoint | AI-based Analytics Tools | D-Board | Intela Workspace"),
    ("Cloud & FinOps (Active Learning)",
     "Cloud Billing Concepts | FinOps Foundation — FOCUS Practitioner (self-study) | "
     "AWS Cost Explorer (familiar) | Azure Cost Management (familiar) | Cloud Cost Tagging & Allocation"),
]

EXPERIENCE = [
    {
        "company": "Deloitte (USI – Offices of the US)",
        "location": "Hyderabad, India",
        "role": "PMO Tax Consultant II",
        "dates": "April 2023 – April 2025",
        "bullets": [
            "Led end-to-end PMO governance for 10+ UK and European client engagements using RAG/RAID frameworks, "
            "milestone trackers, and budget-vs-actual reporting — achieved 100% on-time delivery across direct, "
            "indirect, and statutory streams in a fast-paced, multi-workstream environment.",

            "Standardized cost-tracking, status, and reporting templates on SharePoint + Power BI dashboards — "
            "cut reporting effort 20–25%, shortened leadership review cycles by 2 days/week, and gave executives "
            "single-pane visibility across all engagements.",

            "Coordinated 30+ delivery resources across 4 offices on a centralized tracker, running daily follow-ups "
            "on action items, invoices, approvals, and stakeholder requests — eliminated 15 hrs/week of manual chase "
            "and lifted execution velocity by 40% across time zones.",

            "Owned invoice, PO, SOW, and MSA tracking for the engagement portfolio — reconciled vendor and internal "
            "billing data weekly, flagged variances, and drove resolution with finance and procurement before "
            "month-end close.",

            "Deployed AI-based analytics tools to monitor compliance and cost-risk signals and generate leadership "
            "insights — improved anomaly-detection accuracy by 25% and decision turnaround by 20%.",

            "Owned governance reporting — weekly status, dashboards, RAID logs — for senior leadership and clients "
            "across 3+ time zones, acting as single point of contact between client SMEs and Deloitte teams.",

            "Managed documentation control (SOWs, SOPs, audit files, transition docs) on version-controlled "
            "repositories — passed 100% of internal QA audits.",

            "Mentored 5+ new joiners through structured onboarding, knowledge transfers, and process handovers — "
            "reduced ramp-up time from 6 weeks to 3 weeks.",
        ],
    },
    {
        "company": "Deloitte (USI – Offices of the US)",
        "location": "Hyderabad, India",
        "role": "PMO Tax Consultant I",
        "dates": "June 2022 – April 2023",
        "bullets": [
            "Built compliance-tracking systems and onboarding playbooks for new clients, aligning project timelines "
            "with operational standards and improving client-onboarding efficiency by 30%.",

            "Authored and structured RFPs for MENA and European pursuits with business development — contributed to "
            "4 multinational client wins and 15% revenue uplift in target regions.",

            "Maintained audit-ready documentation repositories (SOWs, SOPs, trackers) on SharePoint with version "
            "control — ensured 100% audit pass rate during transition phases.",
        ],
    },
    {
        "company": "Tooclarity (tooclarity.com)",
        "location": "Hyderabad, India (travel across India)",
        "role": "Project Manager — Institutional Outreach & Onboarding",
        "dates": "2025 (project engagement)",
        "bullets": [
            "Partnered with the founder on go-to-market for institutional onboarding at Tooclarity, an "
            "India-based education-discovery platform — visited colleges and educational institutions in "
            "person, pitched the platform's value proposition to deans and decision-makers, and closed "
            "listing agreements (RFP-style wins at startup scale).",

            "Brought PMO discipline to a founder-led environment: built and ran the outreach pipeline "
            "(prospect lists, meeting cadence, post-visit follow-ups) and a lightweight CRM for "
            "institutional commitments and renewal tracking — kept nothing slipping through the cracks "
            "in a fast-paced, multi-stakeholder setting.",

            "Coordinated between founders, operations, and onboarded institutions during launch — "
            "established weekly review cadence, status updates, and a shared backlog so commitments, "
            "dependencies, and follow-ups stayed visible across the team.",

            "Adapted pitch and process across diverse institution tiers and decision-making styles, "
            "building repeatable onboarding playbooks for newly partnered colleges and supporting the "
            "founder on positioning, pricing conversations, and partner-tier negotiations.",
        ],
    },
    {
        "company": "Wells Fargo",
        "location": "Hyderabad, India",
        "role": "Associate Financial Analyst",
        "dates": "November 2020 – December 2021",
        "bullets": [
            "Audited credits and fee applications on debit, ACH, and cheque flows — surfaced ₹2L+/month in cost "
            "anomalies and mis-applied fees, drove recovery through structured follow-ups with operations and "
            "compliance, and authored controls that prevented repeat occurrence.",

            "Processed Visa debit-card fraud and non-fraud claims (FCM) using bank investigation tools — resolved "
            "40+ cases/day with >98% accuracy and zero compliance breaches.",

            "Performed QA on ATM debit, ACH, and cheque-fraud units — strengthened financial-controls posture and "
            "feedback loops for upstream operations teams.",
        ],
    },
]

MENTORSHIP = [
    "Mentored undergraduate and MBA students across Osmania University and partner colleges on the "
    "college-to-corporate transition — covered skill gaps, professional documentation, stakeholder "
    "communication, and pathways into PMO, consulting, and finance roles.",

    "Ran skill-development sessions bridging classroom learning to corporate expectations — interview "
    "prep, resume craft, workplace etiquette, and basic project / stakeholder skills — preparing "
    "students with corporate-ready capabilities before campus placements.",

    "Advised students through 1:1 conversations on career pathing into PMO, consulting, finance, and "
    "operations roles; built a small, ongoing peer network for follow-up support.",
]

EDUCATION = [
    "MBA, Finance — Osmania University (2018–2020) — GPA 7.3",
    "B.Com — Osmania University (2015–2018) — GPA 7.1",
    "CA Intermediate — ICAI (2018)",
]

CERTIFICATIONS = [
    "PMP Certification — In Progress (target H2 2026)",
    "FinOps Foundation FOCUS Practitioner — Self-study (target Q3 2026)",
    "AWS Cloud Practitioner (CLF-C02) — Planned (target Q3 2026)",
    "NISM (National Institute of Securities Markets): Securities Operations & Risk Management; "
    "Mutual Funds Distributors; Currency Derivatives; Equity Derivatives",
]

AWARDS = [
    "Top Performer — Recognized at Deloitte for 3 quarters (accuracy + leadership).",
    "Fast-track Promotion — Promoted to Tax Consultant II within 10 months.",
    "Mentor & Onboarding Lead — Selected to onboard new hires and lead process handovers.",
]

ADDITIONAL = [
    "Languages: English, Telugu, Hindi",
    "Immediate joiner",
    "Open to: Hyderabad, Bangalore, Mumbai, Remote / Relocation across India",
]

# ---------- COVER LETTER ----------

COVER_LETTER = {
    "to_lines": [
        "Hiring Team",
        "BrightEdge",
        "Remote — India",
    ],
    "date": "25 May 2026",
    "subject": "Application — Cloud Cost Operations / FinOps-PMO Specialist",
    "salutation": "Dear Hiring Team,",
    "paragraphs": [
        "I'm writing to express strong interest in the Cloud Cost Operations role at BrightEdge. Your "
        "JD reads almost like a description of what I've been doing at Deloitte for the past three years — "
        "running trackers, dashboards, invoice and PO coordination, and leadership-ready reporting across "
        "10+ multi-stakeholder engagements — paired with the FinOps direction I've been actively building "
        "toward.",

        "At Deloitte, I owned end-to-end PMO governance for UK and European clients: RAG/RAID frameworks, "
        "Power BI dashboards, weekly status, RAID logs, and budget-vs-actual reporting. I standardized "
        "templates that cut reporting effort 20–25% and shortened leadership review cycles by two days a "
        "week. On follow-ups specifically — the heartbeat of the role you're hiring for — I built a "
        "centralized tracker that eliminated 15 hours a week of manual chase across 30+ resources and "
        "four offices. Invoice, PO, SOW, and MSA reconciliation was a weekly rhythm; variances were "
        "flagged and resolved with finance and procurement before close.",

        "Earlier, at Wells Fargo, I audited fee applications across debit, ACH, and cheque flows and "
        "surfaced ₹2L+/month in cost anomalies and mis-applied fees — work that maps cleanly onto the "
        "cost-anomaly tracking BrightEdge needs, just translated from banking to cloud. I'm "
        "actively closing the cloud gap right now: FinOps Foundation FOCUS Practitioner self-study and "
        "AWS Cloud Practitioner (CLF-C02), both targeted for Q3 2026.",

        "What I bring is the rare bridge profile your team needs: a PMO who already speaks invoices, "
        "variance, and dashboards, with an MBA in Finance and CA Inter behind the numbers, and the "
        "execution discipline to be dependable from day one while ramping into deeper FinOps ownership "
        "over 6–12 months. I'm based in Hyderabad, fully remote-capable, and an immediate joiner.",

        "I'd welcome the chance to walk you through the trackers, dashboards, and follow-up systems "
        "I've built — and discuss how I'd apply that operating rhythm to BrightEdge's cloud-cost world.",
    ],
    "closing": "Warm regards,",
    "signature": "Srikanth Chintala",
    "signature_lines": [
        "+91 81795 98723",
        "srikanthchintala191997@gmail.com",
        "linkedin.com/in/chintalasrikanth97",
    ],
}

# ===================================================================
# DOCX GENERATION
# ===================================================================

def _set_margins(doc: Document, inches: float = 0.6):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    _add_bottom_border(p)


def _add_body(doc: Document, text: str, bold: bool = False, size: int = 10, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def _add_bullet(doc: Document, text: str, size: int = 10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(size)


def _kv_run(p, label: str, value: str):
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)


def build_resume_docx(out_path: Path):
    doc = Document()
    _set_margins(doc, 0.6)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(CONTACT["name"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Headline
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(CONTACT["headline"])
    r.italic = True
    r.font.size = Pt(11)

    # Contact line
    contact_str = f'{CONTACT["location"]} | {CONTACT["phone"]} | {CONTACT["email"]} | {CONTACT["linkedin"]}'
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(contact_str)
    r.font.size = Pt(9)

    # Summary
    _add_section_heading(doc, "Professional Summary")
    _add_body(doc, SUMMARY)

    # Competencies
    _add_section_heading(doc, "Core Competencies")
    for label, value in COMPETENCIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _kv_run(p, f"{label}: ", value)

    # Experience
    _add_section_heading(doc, "Professional Experience")
    for job in EXPERIENCE:
        # Company line
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(job["company"])
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(f'  |  {job["location"]}')
        r2.font.size = Pt(10)

        # Role line
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(job["role"])
        r.italic = True
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(f'    {job["dates"]}')
        r2.italic = True
        r2.font.size = Pt(10)

        for b in job["bullets"]:
            _add_bullet(doc, b)

    # Mentorship & Community
    _add_section_heading(doc, "Mentorship & Community Engagement")
    for m in MENTORSHIP:
        _add_bullet(doc, m)

    # Education
    _add_section_heading(doc, "Education")
    for e in EDUCATION:
        _add_bullet(doc, e)

    # Certifications
    _add_section_heading(doc, "Certifications & Active Learning")
    for c in CERTIFICATIONS:
        _add_bullet(doc, c)

    # Awards
    _add_section_heading(doc, "Awards & Recognition")
    for a in AWARDS:
        _add_bullet(doc, a)

    # Additional
    _add_section_heading(doc, "Additional")
    for a in ADDITIONAL:
        _add_bullet(doc, a)

    doc.save(out_path)


def build_cover_docx(out_path: Path):
    doc = Document()
    _set_margins(doc, 0.8)

    # Sender block (right-aligned name + contact)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r = p.add_run(CONTACT["name"].title())
    r.bold = True
    r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r2 = p2.add_run(f'{CONTACT["location"]}\n{CONTACT["phone"]} | {CONTACT["email"]}\n{CONTACT["linkedin"]}')
    r2.font.size = Pt(10)

    # Date
    _add_body(doc, COVER_LETTER["date"])

    # To
    for line in COVER_LETTER["to_lines"]:
        _add_body(doc, line)

    # Subject
    _add_body(doc, "")
    p = doc.add_paragraph()
    r = p.add_run(f'Subject: {COVER_LETTER["subject"]}')
    r.bold = True
    r.font.size = Pt(11)
    _add_body(doc, "")

    # Salutation
    _add_body(doc, COVER_LETTER["salutation"])
    _add_body(doc, "")

    # Paragraphs
    for para in COVER_LETTER["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(para)
        r.font.size = Pt(11)

    # Closing
    _add_body(doc, "")
    _add_body(doc, COVER_LETTER["closing"])
    p = doc.add_paragraph()
    r = p.add_run(COVER_LETTER["signature"])
    r.bold = True
    r.font.size = Pt(11)
    for line in COVER_LETTER["signature_lines"]:
        _add_body(doc, line, size=10)

    doc.save(out_path)


# ===================================================================
# PDF GENERATION
# ===================================================================

class ResumePDF(FPDF):
    def header(self):
        pass

    def footer(self):
        pass


_PDF_REPLACEMENTS = {
    "\u2014": "-",   # em dash
    "\u2013": "-",   # en dash
    "\u2022": "*",   # bullet (we render our own)
    "\u2018": "'",
    "\u2019": "'",
    "\u201C": '"',
    "\u201D": '"',
    "\u20B9": "INR ",  # rupee
    "\u00A0": " ",
}


def _sanitize_pdf(text: str) -> str:
    for k, v in _PDF_REPLACEMENTS.items():
        text = text.replace(k, v)
    return text


def _pdf_section_heading(pdf: ResumePDF, text: str):
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(31, 58, 95)
    pdf.ln(2)
    pdf.cell(0, 5, text.upper(), new_x="LMARGIN", new_y="NEXT")
    # underline
    y = pdf.get_y()
    pdf.set_draw_color(150, 150, 150)
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(1)


def _pdf_para(pdf: ResumePDF, text: str, font: str = "Helvetica", style: str = "", size: int = 10):
    pdf.set_font(font, style, size)
    pdf.multi_cell(0, 4.5, _sanitize_pdf(text))
    pdf.ln(0.5)


def _pdf_bullet(pdf: ResumePDF, text: str, size: int = 10):
    pdf.set_font("Helvetica", "", size)
    indent = 4
    bullet = "*"
    x_start = pdf.l_margin
    pdf.set_x(x_start)
    pdf.cell(indent, 4.5, bullet)
    pdf.set_x(x_start + indent)
    pdf.multi_cell(pdf.w - pdf.r_margin - (x_start + indent), 4.5, _sanitize_pdf(text))


def _pdf_kv(pdf: ResumePDF, label: str, value: str, size: int = 10):
    pdf.set_font("Helvetica", "B", size)
    pdf.write(4.5, _sanitize_pdf(f"{label}: "))
    pdf.set_font("Helvetica", "", size)
    pdf.write(4.5, _sanitize_pdf(value))
    pdf.ln(5)


def build_resume_pdf(out_path: Path):
    pdf = ResumePDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_margins(left=15, top=12, right=15)
    pdf.add_page()

    # Name
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(31, 58, 95)
    pdf.cell(0, 8, _sanitize_pdf(CONTACT["name"]), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    # Headline
    pdf.set_font("Helvetica", "I", 11)
    pdf.cell(0, 5, _sanitize_pdf(CONTACT["headline"]), align="C", new_x="LMARGIN", new_y="NEXT")

    # Contact
    pdf.set_font("Helvetica", "", 9)
    contact_str = f'{CONTACT["location"]} | {CONTACT["phone"]} | {CONTACT["email"]} | {CONTACT["linkedin"]}'
    pdf.cell(0, 5, _sanitize_pdf(contact_str), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Summary
    _pdf_section_heading(pdf, "Professional Summary")
    _pdf_para(pdf, SUMMARY)

    # Competencies
    _pdf_section_heading(pdf, "Core Competencies")
    for label, value in COMPETENCIES:
        _pdf_kv(pdf, label, value)

    # Experience
    _pdf_section_heading(pdf, "Professional Experience")
    for job in EXPERIENCE:
        pdf.set_font("Helvetica", "B", 11)
        pdf.write(5, _sanitize_pdf(job["company"]))
        pdf.set_font("Helvetica", "", 10)
        pdf.write(5, _sanitize_pdf(f'   |   {job["location"]}'))
        pdf.ln(5)

        pdf.set_font("Helvetica", "BI", 10)
        pdf.write(5, _sanitize_pdf(job["role"]))
        pdf.set_font("Helvetica", "I", 10)
        pdf.write(5, _sanitize_pdf(f'    {job["dates"]}'))
        pdf.ln(5)

        for b in job["bullets"]:
            _pdf_bullet(pdf, b)
        pdf.ln(1)

    # Mentorship & Community
    _pdf_section_heading(pdf, "Mentorship & Community Engagement")
    for m in MENTORSHIP:
        _pdf_bullet(pdf, m)

    # Education
    _pdf_section_heading(pdf, "Education")
    for e in EDUCATION:
        _pdf_bullet(pdf, e)

    # Certifications
    _pdf_section_heading(pdf, "Certifications & Active Learning")
    for c in CERTIFICATIONS:
        _pdf_bullet(pdf, c)

    # Awards
    _pdf_section_heading(pdf, "Awards & Recognition")
    for a in AWARDS:
        _pdf_bullet(pdf, a)

    # Additional
    _pdf_section_heading(pdf, "Additional")
    for a in ADDITIONAL:
        _pdf_bullet(pdf, a)

    pdf.output(str(out_path))


def build_cover_pdf(out_path: Path):
    pdf = ResumePDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(left=20, top=15, right=20)
    pdf.add_page()

    # Sender block right-aligned
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 6, _sanitize_pdf(CONTACT["name"].title()), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 4.5, _sanitize_pdf(CONTACT["location"]), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 4.5, _sanitize_pdf(f'{CONTACT["phone"]} | {CONTACT["email"]}'), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 4.5, _sanitize_pdf(CONTACT["linkedin"]), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Date
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["date"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # To
    for line in COVER_LETTER["to_lines"]:
        pdf.cell(0, 5, _sanitize_pdf(line), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Subject
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 5, _sanitize_pdf(f'Subject: {COVER_LETTER["subject"]}'), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Salutation
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["salutation"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Body
    for para in COVER_LETTER["paragraphs"]:
        pdf.multi_cell(0, 5, _sanitize_pdf(para))
        pdf.ln(2)

    # Closing
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["closing"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["signature"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for line in COVER_LETTER["signature_lines"]:
        pdf.cell(0, 4.5, _sanitize_pdf(line), new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(out_path))


# ===================================================================
# DRIVER
# ===================================================================

def main():
    resume_docx = OUT_DIR / f"{SLUG}.docx"
    resume_pdf = OUT_DIR / f"{SLUG}.pdf"
    cover_docx = OUT_DIR / f"{SLUG}_cover.docx"
    cover_pdf = OUT_DIR / f"{SLUG}_cover.pdf"

    build_resume_docx(resume_docx)
    build_resume_pdf(resume_pdf)
    build_cover_docx(cover_docx)
    build_cover_pdf(cover_pdf)

    print("Generated:")
    for p in (resume_docx, resume_pdf, cover_docx, cover_pdf):
        print(f"  - {p.relative_to(OUT_DIR.parent.parent)}  ({p.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
