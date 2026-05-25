#!/usr/bin/env python3
"""
Generates the tailored Resume and Cover Letter for NES Fircroft Project Coordinator (Diversity).

Outputs:
  - resumes/tailored/NESFircroft_DiversityCoordinator_2026-05-25.docx
  - resumes/tailored/NESFircroft_DiversityCoordinator_2026-05-25.pdf
  - resumes/tailored/NESFircroft_DiversityCoordinator_2026-05-25_cover.docx
  - resumes/tailored/NESFircroft_DiversityCoordinator_2026-05-25_cover.pdf
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable
)


# =================================================================================
# CONTENT — Resume (tailored to JD)
# =================================================================================

CANDIDATE_NAME = "SRIKANTH CHINTALA"
HEADLINE = "Project Coordinator | Consulting Delivery & People-Centric Project Governance"
CONTACT_LINE = "Hyderabad, India 500088 | +91 81795 98723 | srikanthchintala191997@gmail.com"
LINKEDIN = "linkedin.com/in/chintalasrikanth97"

PROFESSIONAL_SUMMARY = (
    "Project Coordinator with 4+ years at Deloitte delivering end-to-end consulting project governance "
    "for UK & European clients across multi-cultural, multi-jurisdiction engagements. Owner of project "
    "trackers, RAID logs, client materials, and workshop coordination — kept 10+ multinational programmes "
    "on schedule with 100% on-time delivery, 25% reporting-effort reduction, and 4 multi-region RFP wins. "
    "Equally strong on stakeholder communications and process improvement; mentored 5+ joiners with "
    "structured onboarding playbooks. Passionate about people-focused, equity-driven project work and "
    "bringing PMO rigour to DEI consulting initiatives. MBA Finance + CA Inter; immediate joiner."
)

CORE_COMPETENCIES = [
    ("Project Coordination & Governance",
     "Project Coordination | Project Plans & Trackers | RAG Status Reporting | RAID Log Management | "
     "Stakeholder Management & Engagement | Scope, Deliverable & Timeline Management | "
     "Workshop & Meeting Coordination | Audit Readiness"),
    ("Financial & Operational",
     "Budget Monitoring | Cost & Variance Analysis | Plan vs. Actual Reporting | "
     "Invoice / PO / SOW / MSA Tracking | Engagement Financial Management | Effort Forecasting"),
    ("Project Delivery",
     "Project Planning & Scheduling | Resource & Capacity Planning | Risk Management | "
     "Cross-functional Collaboration | Process Improvement | Agile / Waterfall | "
     "Implementation Lifecycle | Performance Metrics & KPIs | Vendor Coordination"),
    ("Communication & People",
     "Client-Facing Materials & Presentations | Cross-Cultural Stakeholder Communications | "
     "Mentoring & Onboarding | Cross-Functional Collaboration | Workshop Coordination | "
     "Solution-Focused Problem Solving"),
    ("Tools & Methods",
     "MS Excel (Advanced) | Power BI | MS Project | MS PowerPoint (Client Decks) | Jira | SharePoint | "
     "AI-Based Analytics Tools | D-Board | Intela Workspace | Research & Benchmarking | "
     "Documentation Control"),
]

EXPERIENCE = [
    {
        "company": "Deloitte (USI – Offices of the US)",
        "location": "Hyderabad, India",
        "roles": [
            {
                "title": "PMO Tax Consultant II",
                "dates": "April 2023 – April 2025",
                "bullets": [
                    "Coordinated end-to-end project delivery for 10+ UK & European consulting engagements using RAG/RAID frameworks, milestone trackers, and SharePoint repositories — achieving 100% on-time delivery across multi-jurisdiction streams.",
                    "Standardized project plans, trackers, and reporting templates on SharePoint + Power BI, cutting reporting effort by 20–25% and shortening leadership review cycles by 2 days/week — directly improving processes and ways of working.",
                    "Coordinated 30+ delivery resources across 4 Deloitte offices via a centralized task-tracking system — managing multiple priorities, eliminating 15 hrs/week of manual follow-ups, and lifting cross-functional team efficiency by 40%.",
                    "Mentored 5+ new joiners from diverse academic and regional backgrounds through structured onboarding, knowledge transfers, and process handovers — reducing ramp-up time from 6 weeks to 3 weeks and supporting an inclusive team-integration experience.",
                    "Coordinated client workshops, kickoffs, and review sessions — handling scheduling, material preparation, follow-up actions, and translating discussion points into trackable deliverables for senior stakeholders across 3+ regions.",
                    "Acted as central point of contact for project updates and internal coordination — owning weekly status reports, dashboards, and RAID logs for senior leadership and clients across 3+ time zones and multi-cultural teams, with timely, clear, and professional client communications.",
                    "Deployed AI-based analytics tools to monitor compliance-risk signals and generate leadership insights, improving risk-detection accuracy by 25% and decision turnaround by 20%.",
                    "Managed documentation control (SOWs, SOPs, audit files, transition docs) on version-controlled repositories — passed 100% of internal QA audits.",
                ],
            },
            {
                "title": "PMO Tax Consultant I",
                "dates": "June 2022 – April 2023",
                "bullets": [
                    "Built compliance tracking systems and onboarding playbooks for new clients, aligning project timelines with operational standards and improving client-onboarding efficiency by 30%.",
                    "Conducted research and benchmarking for MENA & European RFP pursuits — preparing client-facing materials, structured insights, and recommendations that contributed to 4 multinational client wins and 15% revenue uplift in target regions.",
                    "Maintained audit-ready documentation repositories (SOWs, SOPs, trackers) on SharePoint with version control, ensuring 100% audit pass rate during transition phases.",
                ],
            },
        ],
    },
    {
        "company": "Wells Fargo",
        "location": "Hyderabad, India",
        "roles": [
            {
                "title": "Associate Financial Analyst",
                "dates": "November 2020 – December 2021",
                "bullets": [
                    "Processed Visa debit-card fraud and non-fraud claims (FCM) using bank investigation tools, resolving 40+ cases/day with >98% accuracy and zero compliance breaches.",
                    "Performed quality reviews and audits on debit/ACH/cheque transactions and customer-account fees, identifying control gaps, strengthening compliance posture, and demonstrating attention to detail and follow-through across high-volume operational reviews.",
                ],
            },
        ],
    },
]

EDUCATION = [
    "MBA, Finance — Osmania University (2018–2020) — GPA 7.3",
    "B.Com — Osmania University (2015–2018) — GPA 7.1",
    "CA Intermediate — ICAI (2018)",
]

CERTIFICATIONS = [
    "NISM (National Institute of Securities Markets): Securities Operations & Risk Management; Mutual Funds Distributors; Currency Derivatives; Equity Derivatives",
    "PMP Certification — In Progress (target H2 2026)",
]

AWARDS = [
    "Top Performer — Recognized at Deloitte for 3 consecutive quarters (delivery accuracy + leadership presence).",
    "Fast-track Promotion — Promoted to Consultant II within 10 months for consistent project-coordination excellence.",
    "Mentor & Onboarding Lead — Selected to onboard new joiners and lead process handovers across cross-cultural delivery teams.",
]

INTERESTS = [
    "Diversity, Equity & Inclusion: Self-driven study of DEI frameworks (APSCo Embrace, Gender Equality Charter), DE&I-in-India reports (BW People, NHRDN), and emerging best-practice in cross-cultural workplace inclusion.",
    "Mentoring: Active mentor for early-career professionals from tier-2 cities transitioning into corporate consulting roles.",
    "Languages: English, Telugu, Hindi — cross-language stakeholder coordination across regional teams.",
]

ADDITIONAL = [
    "Availability: Immediate joiner",
    "Open to: Hyderabad, Bangalore, Mumbai (NES Fircroft offices) | Remote / Hybrid | Relocation across India and UK",
]


# =================================================================================
# CONTENT — Cover Letter
# =================================================================================

COVER_LETTER_DATE = "25 May 2026"
COVER_LETTER_BODY = [
    (
        "Dear Hiring Team at NES Fircroft,"
    ),
    (
        "I'm writing to apply for the Project Coordinator – Diversity role supporting Katy Hall and "
        "the DEI consulting team. The combination of structured project coordination, client-facing "
        "consulting delivery, and people-centred outcomes is exactly the kind of work I want to do "
        "next — and the way NES Fircroft has built DEI into a measurable, industry-shaping practice "
        "(reflected in your APSCo Embrace membership, Gender Equality Charter signatory status, and "
        "Katy's 2026 British Diversity Awards shortlisting) is what made me apply."
    ),
    (
        "Over four years at Deloitte's UK/European tax-consulting practice, I've been the person who "
        "keeps complex, multi-stakeholder programmes on track. I've coordinated 10+ multinational "
        "engagements with 100% on-time delivery, owned RAG/RAID logs and SharePoint trackers across "
        "30+ resources in four offices, prepared client-facing materials and workshop content, and "
        "ran the cadence of weekly status, leadership reviews, and follow-up actions across 3+ time "
        "zones. The bullet I am most proud of, though, is the simplest one: I mentored 5+ new joiners "
        "from diverse academic and regional backgrounds through onboarding, cutting ramp-up from six "
        "weeks to three. That experience — sitting with someone different from me, designing the "
        "process so they could succeed — is the seed of why I'm pursuing this role."
    ),
    (
        "I bring honesty about where I am: I am not coming in with deep DEI domain experience, but I "
        "am coming in with disciplined project coordination, consulting-grade client communications, "
        "and a commitment to learn the DEI craft fast. I've already begun working through the APSCo "
        "Embrace charter, NES Fircroft's DEI thought-leadership posts, and DE&I-in-India industry "
        "reports — and would welcome the chance to discuss how I can apply project rigour to your "
        "DEI programmes from day one."
    ),
    (
        "I'm based in Hyderabad with full availability for Bangalore, Mumbai, hybrid, or relocation. "
        "I'm an immediate joiner. Thank you for considering my application — I'd be glad to discuss "
        "further at your convenience."
    ),
    (
        "Warm regards,\nSrikanth Chintala\n+91 81795 98723 | srikanthchintala191997@gmail.com\n"
        "linkedin.com/in/chintalasrikanth97"
    ),
]


# =================================================================================
# .DOCX GENERATION
# =================================================================================

def style_paragraph(p, *, size=10.5, bold=False, italic=False, color=None,
                    align=None, space_after=None, space_before=None):
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    for run in p.runs:
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x14, 0x3B, 0x6E)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # underline via bottom border
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '143B6E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_resume_docx(out_path: Path):
    doc = Document()

    # Set page margins (narrow)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Header — name
    p = doc.add_paragraph()
    run = p.add_run(CANDIDATE_NAME)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x14, 0x3B, 0x6E)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    # Headline
    p = doc.add_paragraph()
    run = p.add_run(HEADLINE)
    run.italic = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    # Contact
    p = doc.add_paragraph()
    run = p.add_run(CONTACT_LINE)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    # LinkedIn
    p = doc.add_paragraph()
    run = p.add_run(LINKEDIN)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # Professional summary
    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph(PROFESSIONAL_SUMMARY)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10.5)

    # Core competencies
    add_section_heading(doc, "Core Competencies")
    for category, items in CORE_COMPETENCIES:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{category}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(items)
        r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)

    # Professional experience
    add_section_heading(doc, "Professional Experience")
    for exp in EXPERIENCE:
        # Company line
        p = doc.add_paragraph()
        run = p.add_run(f"{exp['company']}  |  {exp['location']}")
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(4)

        for role in exp["roles"]:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{role['title']}")
            r1.italic = True
            r1.font.size = Pt(10.5)
            r1.bold = True
            r2 = p.add_run(f"   {role['dates']}")
            r2.font.size = Pt(10)
            r2.italic = True
            p.paragraph_format.space_after = Pt(2)

            for bullet in role["bullets"]:
                bp = doc.add_paragraph(bullet, style='List Bullet')
                for run in bp.runs:
                    run.font.size = Pt(10.5)
                bp.paragraph_format.space_after = Pt(0)

    # Education
    add_section_heading(doc, "Education")
    for ed in EDUCATION:
        bp = doc.add_paragraph(ed, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(10.5)
        bp.paragraph_format.space_after = Pt(0)

    # Certifications
    add_section_heading(doc, "Certifications")
    for c in CERTIFICATIONS:
        bp = doc.add_paragraph(c, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(10.5)
        bp.paragraph_format.space_after = Pt(0)

    # Awards
    add_section_heading(doc, "Awards & Recognition")
    for a in AWARDS:
        bp = doc.add_paragraph(a, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(10.5)
        bp.paragraph_format.space_after = Pt(0)

    # Interests
    add_section_heading(doc, "Interests & Continuous Learning")
    for i in INTERESTS:
        bp = doc.add_paragraph(i, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(10.5)
        bp.paragraph_format.space_after = Pt(0)

    # Additional
    add_section_heading(doc, "Additional")
    for a in ADDITIONAL:
        bp = doc.add_paragraph(a, style='List Bullet')
        for run in bp.runs:
            run.font.size = Pt(10.5)
        bp.paragraph_format.space_after = Pt(0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] Resume DOCX → {out_path}")


def build_cover_letter_docx(out_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Header — sender info
    for line in [CANDIDATE_NAME, "Hyderabad, India 500088", "+91 81795 98723",
                 "srikanthchintala191997@gmail.com", LINKEDIN]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        if line == CANDIDATE_NAME:
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x14, 0x3B, 0x6E)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Date
    p = doc.add_paragraph()
    p.add_run(COVER_LETTER_DATE).font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)

    # Recipient
    for line in ["Hiring Team", "NES Fircroft", "Re: Project Coordinator – Diversity"]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(11)
        if line.startswith("Re:"):
            r.bold = True
        p.paragraph_format.space_after = Pt(0)

    # Body
    for para in COVER_LETTER_BODY:
        p = doc.add_paragraph()
        for line in para.split("\n"):
            r = p.add_run(line)
            r.font.size = Pt(11)
            r.add_break()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] Cover Letter DOCX → {out_path}")


# =================================================================================
# .PDF GENERATION (via ReportLab — direct, no LibreOffice required)
# =================================================================================

def build_resume_pdf(out_path: Path):
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        title=f"{CANDIDATE_NAME} — Resume — NES Fircroft",
    )

    styles = getSampleStyleSheet()
    deep = HexColor("#143B6E")

    name_style = ParagraphStyle(
        "Name", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=20, textColor=deep, alignment=TA_CENTER, leading=22, spaceAfter=2,
    )
    headline_style = ParagraphStyle(
        "Headline", parent=styles["Normal"], fontName="Helvetica-Oblique",
        fontSize=11, alignment=TA_CENTER, leading=13, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10, alignment=TA_CENTER, leading=12, spaceAfter=0,
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=11, textColor=deep, leading=14, spaceBefore=8, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10, alignment=TA_JUSTIFY, leading=12.5, spaceAfter=3,
    )
    role_company_style = ParagraphStyle(
        "Company", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=11, leading=13, spaceBefore=4, spaceAfter=0,
    )
    role_title_style = ParagraphStyle(
        "RoleTitle", parent=styles["Normal"], fontName="Helvetica-Oblique",
        fontSize=10.5, leading=12.5, spaceAfter=2,
    )

    story = []

    # Header
    story.append(Paragraph(CANDIDATE_NAME, name_style))
    story.append(Paragraph(HEADLINE, headline_style))
    story.append(Paragraph(CONTACT_LINE, contact_style))
    story.append(Paragraph(LINKEDIN, contact_style))

    def add_section(title):
        story.append(Spacer(1, 4))
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.6, color=deep, spaceBefore=0, spaceAfter=2))

    def bullets(items):
        story.append(ListFlowable(
            [ListItem(Paragraph(i, body_style), leftIndent=12) for i in items],
            bulletType='bullet', start='•', leftIndent=14,
        ))

    # Sections
    add_section("Professional Summary")
    story.append(Paragraph(PROFESSIONAL_SUMMARY, body_style))

    add_section("Core Competencies")
    for cat, items in CORE_COMPETENCIES:
        story.append(Paragraph(f"<b>{cat}:</b> {items}", body_style))

    add_section("Professional Experience")
    for exp in EXPERIENCE:
        story.append(Paragraph(f"{exp['company']}  |  {exp['location']}", role_company_style))
        for role in exp["roles"]:
            story.append(Paragraph(
                f"<b>{role['title']}</b> &nbsp;&nbsp; <i>{role['dates']}</i>", role_title_style
            ))
            bullets(role["bullets"])

    add_section("Education")
    bullets(EDUCATION)

    add_section("Certifications")
    bullets(CERTIFICATIONS)

    add_section("Awards & Recognition")
    bullets(AWARDS)

    add_section("Interests & Continuous Learning")
    bullets(INTERESTS)

    add_section("Additional")
    bullets(ADDITIONAL)

    doc.build(story)
    print(f"[OK] Resume PDF  → {out_path}")


def build_cover_letter_pdf(out_path: Path):
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        title=f"{CANDIDATE_NAME} — Cover Letter — NES Fircroft",
    )

    styles = getSampleStyleSheet()
    deep = HexColor("#143B6E")

    name_style = ParagraphStyle(
        "Name", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=14, textColor=deep, leading=16, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, leading=13, spaceAfter=0,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, alignment=TA_JUSTIFY, leading=14, spaceAfter=8, spaceBefore=4,
    )

    story = []
    story.append(Paragraph(CANDIDATE_NAME, name_style))
    for line in ["Hyderabad, India 500088", "+91 81795 98723",
                 "srikanthchintala191997@gmail.com", LINKEDIN]:
        story.append(Paragraph(line, contact_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph(COVER_LETTER_DATE, contact_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Hiring Team", contact_style))
    story.append(Paragraph("NES Fircroft", contact_style))
    story.append(Paragraph("<b>Re: Project Coordinator – Diversity</b>", contact_style))
    story.append(Spacer(1, 8))

    for para in COVER_LETTER_BODY:
        # Convert any newlines into <br/> for ReportLab
        para_html = para.replace("\n", "<br/>")
        story.append(Paragraph(para_html, body_style))

    doc.build(story)
    print(f"[OK] Cover Letter PDF → {out_path}")


# =================================================================================
# MAIN
# =================================================================================

def main():
    base = Path(__file__).resolve().parents[1]  # f/
    tailored = base / "resumes" / "tailored"
    slug = "NESFircroft_DiversityCoordinator_2026-05-25"

    build_resume_docx(tailored / f"{slug}.docx")
    build_resume_pdf(tailored / f"{slug}.pdf")
    build_cover_letter_docx(tailored / f"{slug}_cover.docx")
    build_cover_letter_pdf(tailored / f"{slug}_cover.pdf")


if __name__ == "__main__":
    main()
