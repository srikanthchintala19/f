"""
Generates the tailored resume + cover letter for Citrin Cooperman India LLP
Project Manager role in both .docx and .pdf.

Run: python3 scripts/generate_citrincooperman_docs.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

OUT_DIR = Path(__file__).resolve().parent.parent / "resumes" / "tailored"
OUT_DIR.mkdir(parents=True, exist_ok=True)
SLUG = "citrincooperman_project-manager_2026-05-25"

# ---------- CONTENT ----------

CONTACT = {
    "name": "SRIKANTH CHINTALA",
    "headline": "Project Manager | PMO, Finance & MIS Reporting | Cross-Service-Line Change Agent",
    "location": "Hyderabad, India 500088",
    "phone": "+91 81795 98723",
    "email": "srikanthchintala191997@gmail.com",
    "linkedin": "linkedin.com/in/chintalasrikanth97",
}

SUMMARY = (
    "Project Manager and PMO professional with 4+ years at Deloitte (USI – Offices of the US) running "
    "end-to-end project governance, MIS reporting, and stakeholder engagement across multiple Service "
    "Lines and Business Units for UK and European clients. Built and ran the RAG/RAID, dashboard, "
    "budget-vs-actual, and leadership-MIS machinery that kept 10+ multi-stakeholder, time-zone-spanning "
    "engagements on time and audit-ready — turning messy delivery into clean, partner-ready trackers. "
    "Quantified change-agent record: 100% on-time delivery, 25% MIS-effort reduction, 15 hrs/week of "
    "follow-ups eliminated, 40% velocity lift, and 4 RFP wins. Comfortable communicating analyst-to-"
    "partner. Strong on Excel, Power BI, and PowerPoint; actively building SQL and Tableau proficiency. "
    "MBA Finance + CA Inter; immediate joiner; based in Hyderabad."
)

COMPETENCIES = [
    ("Project Management & MIS Reporting",
     "Project Management Office (PMO) | Cross-Service-Line / Multi-BU Coordination | "
     "Leadership MIS & Status Reporting | RAG / RAID Frameworks | Milestone & Dependency Tracking | "
     "Audit Readiness | Documentation Control"),
    ("Finance & Operations",
     "Financial Tracking | Budget vs Actual | Variance & Plan-vs-Actual Analysis | Forecasting | "
     "Invoice / PO / SOW / MSA Tracking | Vendor Coordination | Operational Reviews | Process Improvement"),
    ("Data, BI & Analytics",
     "Power BI (Advanced) | Tableau (Familiar — Power BI native; equivalent semantic model & visual "
     "grammar) | SQL (Active Learning) | Advanced MS Excel — Pivot Tables, XLOOKUP / INDEX-MATCH, SUMIFS, "
     "Power Query, Conditional Formatting, Data Validation, Basic Macros | AI-based Analytics Tools"),
    ("Communication & Stakeholder Engagement",
     "Analyst-to-Partner Communication | Multi-Time-Zone Coordination (UK / Europe / India) | "
     "Client SPOC | Executive & Leadership Decks (PowerPoint) | Cross-functional Collaboration | "
     "Mentorship & Onboarding"),
    ("Tools & Platforms",
     "MS PowerPoint | MS Project | SharePoint | Jira | D-Board | Intela Workspace"),
]

EXPERIENCE = [
    {
        "company": "Deloitte (USI – Offices of the US)",
        "location": "Hyderabad, India",
        "role": "PMO Tax Consultant II",
        "dates": "April 2023 – April 2025",
        "bullets": [
            "Led end-to-end PMO and project management for 10+ UK and European client engagements across "
            "multiple Service Lines (direct, indirect, statutory) — using RAG/RAID frameworks, milestone "
            "trackers, and budget-vs-actual reporting — achieved 100% on-time delivery in a fast-paced, "
            "multi-workstream environment.",

            "Standardized MIS reporting, status, and dashboard templates on SharePoint + Power BI for "
            "senior leadership — cut MIS preparation effort 20–25%, shortened review cycles by 2 days/week, "
            "and gave partners single-pane visibility across the engagement portfolio.",

            "Coordinated 30+ delivery resources across 4 offices on a centralized tracker, running daily "
            "follow-ups on action items, approvals, and stakeholder requests — eliminated 15 hrs/week of "
            "manual chase and lifted execution velocity by 40% across time zones.",

            "Owned invoice, PO, SOW, and MSA tracking and reconciliation for the engagement portfolio — "
            "flagged variances, drove resolution with finance and procurement, and kept Service Lines "
            "aligned to monthly close.",

            "Championed AI-based analytics adoption to monitor compliance and risk signals, surface "
            "anomalies, and accelerate leadership decisions — improved anomaly-detection accuracy by 25% "
            "and decision turnaround by 20%; positioned as a thoughtful change agent within the portfolio.",

            "Owned governance reporting — weekly status, dashboards, RAID logs, leadership decks — for "
            "senior leadership and clients across 3+ time zones; acted as single point of contact between "
            "client SMEs and Deloitte teams, adapting communication from analyst-level detail to "
            "partner-level brevity.",

            "Managed documentation control (SOWs, SOPs, audit files, transition docs) on version-controlled "
            "repositories — passed 100% of internal QA audits.",

            "Mentored 5+ new joiners through structured onboarding, knowledge transfers, and process "
            "handovers — reduced ramp-up time from 6 weeks to 3 weeks.",
        ],
    },
    {
        "company": "Deloitte (USI – Offices of the US)",
        "location": "Hyderabad, India",
        "role": "PMO Tax Consultant I",
        "dates": "June 2022 – April 2023",
        "bullets": [
            "Built compliance-tracking systems and onboarding playbooks for new clients, aligning project "
            "timelines with operational standards and improving client-onboarding efficiency by 30%.",

            "Authored and structured RFPs for MENA and European pursuits with business development — "
            "contributed to 4 multinational client wins and 15% revenue uplift in target regions.",

            "Maintained audit-ready documentation repositories (SOWs, SOPs, trackers) on SharePoint with "
            "version control — ensured 100% audit pass rate during transition phases.",
        ],
    },
    {
        "company": "Tooclarity (tooclarity.com)",
        "location": "Hyderabad, India (travel across India)",
        "role": "Project Manager — Institutional Outreach & Onboarding",
        "dates": "2025 (project engagement)",
        "bullets": [
            "Partnered with the founder on go-to-market for institutional onboarding at Tooclarity, an "
            "India-based education-discovery platform — visited colleges and educational institutions in "
            "person, pitched the platform's value proposition to deans and decision-makers, and closed "
            "listing agreements (RFP-style wins at startup scale).",

            "Brought PMO discipline to a founder-led environment: built and ran the outreach pipeline "
            "(prospect lists, meeting cadence, post-visit follow-ups) and a lightweight CRM for "
            "institutional commitments and renewal tracking — kept nothing slipping through the cracks "
            "in a fast-paced, multi-stakeholder setting.",

            "Coordinated between founders, operations, and onboarded institutions during launch — "
            "established weekly review cadence, status updates, and a shared backlog so commitments, "
            "dependencies, and follow-ups stayed visible across the team.",

            "Adapted pitch and process across diverse institution tiers and decision-making styles, "
            "building repeatable onboarding playbooks for newly partnered colleges.",
        ],
    },
    {
        "company": "Wells Fargo",
        "location": "Hyderabad, India",
        "role": "Associate Financial Analyst",
        "dates": "November 2020 – December 2021",
        "bullets": [
            "Audited credits and fee applications on debit, ACH, and cheque flows — surfaced ₹2L+/month "
            "in mis-applied fees, drove recovery through structured follow-ups with operations and "
            "compliance, and authored controls that prevented repeat occurrence.",

            "Processed Visa debit-card fraud and non-fraud claims (FCM) using bank investigation tools — "
            "resolved 40+ cases/day with >98% accuracy and zero compliance breaches.",

            "Performed QA on ATM debit, ACH, and cheque-fraud units — strengthened financial-controls "
            "posture and feedback loops for upstream operations teams.",
        ],
    },
]

MENTORSHIP = [
    "Mentored undergraduate and MBA students across Osmania University and partner colleges on the "
    "college-to-corporate transition — covered skill gaps, professional documentation, stakeholder "
    "communication, and pathways into PMO, consulting, and finance roles.",

    "Ran skill-development sessions bridging classroom learning to corporate expectations — interview "
    "prep, resume craft, and basic project / stakeholder skills — preparing students with corporate-ready "
    "capabilities before campus placements.",
]

EDUCATION = [
    "MBA, Finance — Osmania University (2018–2020) — GPA 7.3",
    "B.Com — Osmania University (2015–2018) — GPA 7.1",
    "CA Intermediate — ICAI (2018)",
]

CERTIFICATIONS = [
    "PMP Certification — In Progress (target H2 2026)",
    "SQL Fundamentals — Active learning (target Q3 2026)",
    "Tableau Desktop Specialist — Planned (target Q3 2026)",
    "NISM (National Institute of Securities Markets): Securities Operations & Risk Management; "
    "Mutual Funds Distributors; Currency Derivatives; Equity Derivatives",
]

AWARDS = [
    "Top Performer — Recognised at Deloitte for 3 quarters (accuracy + leadership).",
    "Fast-track Promotion — Promoted to Tax Consultant II within 10 months.",
    "Mentor & Onboarding Lead — Selected to onboard new hires and lead process handovers.",
]

ADDITIONAL = [
    "Languages: English, Telugu, Hindi",
    "Immediate joiner",
    "Open to: Hyderabad (preferred), Bangalore, Mumbai, Remote / Relocation across India",
]

# ---------- COVER LETTER ----------

COVER_LETTER = {
    "to_lines": [
        "Hiring Team",
        "Citrin Cooperman India LLP",
        "Hyderabad",
    ],
    "date": "25 May 2026",
    "subject": "Application — Project Manager",
    "salutation": "Dear Hiring Team,",
    "paragraphs": [
        "I'm writing to apply for the Project Manager role at Citrin Cooperman India LLP. The JD reads "
        "almost like a description of the work I've been doing for the last four years at Deloitte (USI - "
        "Offices of the US): cross-Service-Line PMO, MIS and dashboard reporting, partner-ready "
        "communication, and a steady record of challenging existing processes to find more scalable ways "
        "of running a growing engagement portfolio.",

        "At Deloitte I owned end-to-end PMO governance for 10+ UK and European client engagements across "
        "direct, indirect, and statutory Service Lines. I standardized MIS reporting, status, and "
        "dashboard templates on SharePoint and Power BI for senior leadership — cutting MIS preparation "
        "effort 20-25%, shortening leadership review cycles by two days a week, and giving partners "
        "single-pane visibility across the portfolio. On the change-agent front specifically: I built a "
        "centralized follow-up tracker that eliminated 15 hours a week of manual chase across 30+ "
        "resources and four offices, championed AI-based analytics adoption that lifted anomaly-detection "
        "accuracy by 25% and decision turnaround by 20%, and contributed to 4 RFP wins authoring "
        "structured proposals for MENA and European pursuits.",

        "More recently I supported Tooclarity, an India-based education platform, partnering with the "
        "founder on in-person institutional outreach across diverse college tiers. I built the outreach "
        "pipeline, ran follow-ups and a lightweight CRM, and adapted the pitch and onboarding playbooks "
        "for tier-1, tier-2 and regional institutions — bringing the same project rigour I used at "
        "Deloitte to a fast-paced, founder-led environment. Alongside that, I've been mentoring "
        "undergraduate and MBA students at Osmania University and partner colleges on the college-to-"
        "corporate transition — interview prep, resume craft, and entry pathways into PMO, consulting "
        "and finance roles.",

        "I'm transparent about two technical gaps versus the JD. SQL and Tableau are not currently on my "
        "resume — I'm Power BI native (and Power BI's semantic model and visual grammar transfer cleanly "
        "to Tableau), and I'm actively closing both gaps with formal targets in Q3 2026: SQL "
        "Fundamentals and Tableau Desktop Specialist. I'm happy to demonstrate a Power BI dashboard in "
        "interview, and to rebuild it in Tableau within a 1-2 week ramp.",

        "What I bring is a relatively rare combination for this role: an MBA in Finance and CA "
        "Intermediate behind 4+ years of PMO at a top-tier US-firm ecosystem (Deloitte USI), comfort "
        "communicating from analyst to partner, and a quantified record of process improvement that "
        "matches the change-agent posture you describe. Citrin Cooperman's recent Hyderabad expansion "
        "and India growth strategy fits cleanly with where I am — Hyderabad-based, fully remote-capable, "
        "and an immediate joiner.",

        "I'd welcome the chance to walk you through the trackers, dashboards, and MIS systems I've built "
        "at Deloitte and Tooclarity, and discuss how I'd apply that operating rhythm to Citrin "
        "Cooperman's multi-Service-Line engagements.",
    ],
    "closing": "Warm regards,",
    "signature": "Srikanth Chintala",
    "signature_lines": [
        "+91 81795 98723",
        "srikanthchintala191997@gmail.com",
        "linkedin.com/in/chintalasrikanth97",
    ],
}

# ===================================================================
# DOCX HELPERS
# ===================================================================

def _set_margins(doc: Document, inches: float = 0.6):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    _add_bottom_border(p)


def _add_body(doc: Document, text: str, bold: bool = False, size: int = 10, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def _add_bullet(doc: Document, text: str, size: int = 10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.size = Pt(size)


def _kv_run(p, label: str, value: str):
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.size = Pt(10)


def build_resume_docx(out_path: Path):
    doc = Document()
    _set_margins(doc, 0.6)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(CONTACT["name"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(CONTACT["headline"])
    r.italic = True
    r.font.size = Pt(11)

    contact_str = f'{CONTACT["location"]} | {CONTACT["phone"]} | {CONTACT["email"]} | {CONTACT["linkedin"]}'
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(contact_str)
    r.font.size = Pt(9)

    _add_section_heading(doc, "Professional Summary")
    _add_body(doc, SUMMARY)

    _add_section_heading(doc, "Core Competencies")
    for label, value in COMPETENCIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _kv_run(p, f"{label}: ", value)

    _add_section_heading(doc, "Professional Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(job["company"])
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(f'  |  {job["location"]}')
        r2.font.size = Pt(10)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(job["role"])
        r.italic = True
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(f'    {job["dates"]}')
        r2.italic = True
        r2.font.size = Pt(10)

        for b in job["bullets"]:
            _add_bullet(doc, b)

    _add_section_heading(doc, "Mentorship & Community Engagement")
    for m in MENTORSHIP:
        _add_bullet(doc, m)

    _add_section_heading(doc, "Education")
    for e in EDUCATION:
        _add_bullet(doc, e)

    _add_section_heading(doc, "Certifications & Active Learning")
    for c in CERTIFICATIONS:
        _add_bullet(doc, c)

    _add_section_heading(doc, "Awards & Recognition")
    for a in AWARDS:
        _add_bullet(doc, a)

    _add_section_heading(doc, "Additional")
    for a in ADDITIONAL:
        _add_bullet(doc, a)

    doc.save(out_path)


def build_cover_docx(out_path: Path):
    doc = Document()
    _set_margins(doc, 0.8)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r = p.add_run(CONTACT["name"].title())
    r.bold = True
    r.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r2 = p2.add_run(f'{CONTACT["location"]}\n{CONTACT["phone"]} | {CONTACT["email"]}\n{CONTACT["linkedin"]}')
    r2.font.size = Pt(10)

    _add_body(doc, COVER_LETTER["date"])

    for line in COVER_LETTER["to_lines"]:
        _add_body(doc, line)

    _add_body(doc, "")
    p = doc.add_paragraph()
    r = p.add_run(f'Subject: {COVER_LETTER["subject"]}')
    r.bold = True
    r.font.size = Pt(11)
    _add_body(doc, "")

    _add_body(doc, COVER_LETTER["salutation"])
    _add_body(doc, "")

    for para in COVER_LETTER["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(para)
        r.font.size = Pt(11)

    _add_body(doc, "")
    _add_body(doc, COVER_LETTER["closing"])
    p = doc.add_paragraph()
    r = p.add_run(COVER_LETTER["signature"])
    r.bold = True
    r.font.size = Pt(11)
    for line in COVER_LETTER["signature_lines"]:
        _add_body(doc, line, size=10)

    doc.save(out_path)


# ===================================================================
# PDF
# ===================================================================

class ResumePDF(FPDF):
    def header(self): pass
    def footer(self): pass


_PDF_REPLACEMENTS = {
    "\u2014": "-", "\u2013": "-", "\u2022": "*",
    "\u2018": "'", "\u2019": "'", "\u201C": '"', "\u201D": '"',
    "\u20B9": "INR ", "\u00A0": " ",
}

def _sanitize_pdf(text: str) -> str:
    for k, v in _PDF_REPLACEMENTS.items():
        text = text.replace(k, v)
    return text


def _pdf_section_heading(pdf, text: str):
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(31, 58, 95)
    pdf.ln(2)
    pdf.cell(0, 5, text.upper(), new_x="LMARGIN", new_y="NEXT")
    y = pdf.get_y()
    pdf.set_draw_color(150, 150, 150)
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(1)


def _pdf_para(pdf, text, size=10):
    pdf.set_font("Helvetica", "", size)
    pdf.multi_cell(0, 4.5, _sanitize_pdf(text))
    pdf.ln(0.5)


def _pdf_bullet(pdf, text, size=10):
    pdf.set_font("Helvetica", "", size)
    indent = 4
    x_start = pdf.l_margin
    pdf.set_x(x_start)
    pdf.cell(indent, 4.5, "*")
    pdf.set_x(x_start + indent)
    pdf.multi_cell(pdf.w - pdf.r_margin - (x_start + indent), 4.5, _sanitize_pdf(text))


def _pdf_kv(pdf, label, value, size=10):
    pdf.set_font("Helvetica", "B", size)
    pdf.write(4.5, _sanitize_pdf(f"{label}: "))
    pdf.set_font("Helvetica", "", size)
    pdf.write(4.5, _sanitize_pdf(value))
    pdf.ln(5)


def build_resume_pdf(out_path: Path):
    pdf = ResumePDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_margins(left=15, top=12, right=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(31, 58, 95)
    pdf.cell(0, 8, _sanitize_pdf(CONTACT["name"]), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Helvetica", "I", 11)
    pdf.cell(0, 5, _sanitize_pdf(CONTACT["headline"]), align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 9)
    contact_str = f'{CONTACT["location"]} | {CONTACT["phone"]} | {CONTACT["email"]} | {CONTACT["linkedin"]}'
    pdf.cell(0, 5, _sanitize_pdf(contact_str), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    _pdf_section_heading(pdf, "Professional Summary")
    _pdf_para(pdf, SUMMARY)

    _pdf_section_heading(pdf, "Core Competencies")
    for label, value in COMPETENCIES:
        _pdf_kv(pdf, label, value)

    _pdf_section_heading(pdf, "Professional Experience")
    for job in EXPERIENCE:
        pdf.set_font("Helvetica", "B", 11)
        pdf.write(5, _sanitize_pdf(job["company"]))
        pdf.set_font("Helvetica", "", 10)
        pdf.write(5, _sanitize_pdf(f'   |   {job["location"]}'))
        pdf.ln(5)

        pdf.set_font("Helvetica", "BI", 10)
        pdf.write(5, _sanitize_pdf(job["role"]))
        pdf.set_font("Helvetica", "I", 10)
        pdf.write(5, _sanitize_pdf(f'    {job["dates"]}'))
        pdf.ln(5)

        for b in job["bullets"]:
            _pdf_bullet(pdf, b)
        pdf.ln(1)

    _pdf_section_heading(pdf, "Mentorship & Community Engagement")
    for m in MENTORSHIP:
        _pdf_bullet(pdf, m)

    _pdf_section_heading(pdf, "Education")
    for e in EDUCATION:
        _pdf_bullet(pdf, e)

    _pdf_section_heading(pdf, "Certifications & Active Learning")
    for c in CERTIFICATIONS:
        _pdf_bullet(pdf, c)

    _pdf_section_heading(pdf, "Awards & Recognition")
    for a in AWARDS:
        _pdf_bullet(pdf, a)

    _pdf_section_heading(pdf, "Additional")
    for a in ADDITIONAL:
        _pdf_bullet(pdf, a)

    pdf.output(str(out_path))


def build_cover_pdf(out_path: Path):
    pdf = ResumePDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(left=20, top=15, right=20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 6, _sanitize_pdf(CONTACT["name"].title()), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 4.5, _sanitize_pdf(CONTACT["location"]), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 4.5, _sanitize_pdf(f'{CONTACT["phone"]} | {CONTACT["email"]}'), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 4.5, _sanitize_pdf(CONTACT["linkedin"]), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["date"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    for line in COVER_LETTER["to_lines"]:
        pdf.cell(0, 5, _sanitize_pdf(line), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 5, _sanitize_pdf(f'Subject: {COVER_LETTER["subject"]}'), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["salutation"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    for para in COVER_LETTER["paragraphs"]:
        pdf.multi_cell(0, 5, _sanitize_pdf(para))
        pdf.ln(2)

    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["closing"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 5, _sanitize_pdf(COVER_LETTER["signature"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for line in COVER_LETTER["signature_lines"]:
        pdf.cell(0, 4.5, _sanitize_pdf(line), new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(out_path))


def main():
    rdocx = OUT_DIR / f"{SLUG}.docx"
    rpdf  = OUT_DIR / f"{SLUG}.pdf"
    cdocx = OUT_DIR / f"{SLUG}_cover.docx"
    cpdf  = OUT_DIR / f"{SLUG}_cover.pdf"

    build_resume_docx(rdocx)
    build_resume_pdf(rpdf)
    build_cover_docx(cdocx)
    build_cover_pdf(cpdf)

    print("Generated:")
    for p in (rdocx, rpdf, cdocx, cpdf):
        print(f"  - {p.relative_to(OUT_DIR.parent.parent)}  ({p.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
